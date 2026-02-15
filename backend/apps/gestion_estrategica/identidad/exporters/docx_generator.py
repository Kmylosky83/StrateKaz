"""
Generador de DOCX para Identidad Corporativa v4.0
==================================================

Genera documentos Word editables con la identidad corporativa
usando python-docx para maxima compatibilidad.

Caracteristicas:
- Documento editable Microsoft Word
- Estilos profesionales predefinidos
- Colores dinamicos desde TenantBranding
- Tablas formateadas
- Header/footer con branding

NOTA v4.0: Metodos de politica eliminados.
Las politicas se gestionan desde Gestion Documental (tipo_documento=POL).
"""

import os
from io import BytesIO
from datetime import datetime

# python-docx para generacion de documentos Word
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def hex_to_rgb_color(hex_color: str):
    """
    Convierte color hexadecimal a RGBColor de python-docx.

    Args:
        hex_color: Color en formato '#RRGGBB' o 'RRGGBB'

    Returns:
        RGBColor: Objeto de color para python-docx
    """
    if not DOCX_AVAILABLE:
        return None

    # Remover # si existe
    hex_color = hex_color.lstrip('#')

    # Parsear componentes RGB
    try:
        r = int(hex_color[0:2], 16)
        g = int(hex_color[2:4], 16)
        b = int(hex_color[4:6], 16)
        return RGBColor(r, g, b)
    except (ValueError, IndexError):
        # Color por defecto si falla el parseo
        return RGBColor(0x2d, 0x37, 0x48)


class IdentidadDOCXGenerator:
    """
    Generador de documentos DOCX para Identidad Corporativa.

    Soporta:
    - Documento completo de identidad corporativa
    - Valores Corporativos
    - Alcances del Sistema

    Los colores se obtienen dinamicamente desde TenantBranding.

    NOTA v4.0: Las politicas se gestionan desde Gestion Documental.
    """

    def __init__(self, empresa=None):
        """
        Inicializa el generador.

        Args:
            empresa: Instancia de EmpresaConfig (opcional)
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx no esta instalado. Ejecute: pip install python-docx")

        self.empresa = empresa
        self.logo_path = None
        self.colors = self._load_colors()

        if empresa and empresa.logo:
            self._load_logo()

    def _load_colors(self):
        """
        Carga colores dinamicos desde TenantBranding o usa defaults.

        Returns:
            dict: Diccionario con colores RGBColor
        """
        # Colores por defecto
        defaults = {
            'primary': '#1a365d',       # Azul oscuro
            'secondary': '#3182ce',     # Azul medio
            'text': '#2d3748',          # Gris oscuro
            'muted': '#718096',         # Gris claro
            'success': '#48bb78',       # Verde
            'warning': '#ed8a36',       # Naranja
        }

        # Intentar cargar colores desde TenantBranding
        if self.empresa:
            try:
                from apps.core.models import TenantBranding
                branding = TenantBranding.objects.filter(
                    empresa=self.empresa,
                    is_active=True
                ).first()

                if branding:
                    defaults['primary'] = branding.primary_color or defaults['primary']
                    defaults['secondary'] = branding.secondary_color or defaults['secondary']
                    if hasattr(branding, 'accent_color') and branding.accent_color:
                        defaults['success'] = branding.accent_color
            except Exception:
                pass  # Usar defaults si hay error

        # Convertir hex a RGBColor
        return {
            key: hex_to_rgb_color(value)
            for key, value in defaults.items()
        }

    def _load_logo(self):
        """Carga la ruta del logo de la empresa"""
        try:
            if self.empresa.logo and hasattr(self.empresa.logo, 'path'):
                if os.path.exists(self.empresa.logo.path):
                    self.logo_path = self.empresa.logo.path
        except Exception:
            pass

    def _get_empresa_info(self):
        """Obtiene informacion de la empresa"""
        if self.empresa:
            return {
                'razon_social': self.empresa.razon_social,
                'nit': self.empresa.nit,
                'nombre_comercial': getattr(self.empresa, 'nombre_comercial', ''),
            }
        return {
            'razon_social': 'Empresa',
            'nit': '',
            'nombre_comercial': '',
        }

    def _setup_styles(self, document):
        """Configura estilos del documento"""
        styles = document.styles

        # Estilo para titulos principales
        if 'CustomHeading1' not in [s.name for s in styles]:
            style = styles.add_style('CustomHeading1', WD_STYLE_TYPE.PARAGRAPH)
            style.font.name = 'Arial'
            style.font.size = Pt(18)
            style.font.bold = True
            style.font.color.rgb = self.colors['primary']
            style.paragraph_format.space_after = Pt(12)
            style.paragraph_format.space_before = Pt(18)

        # Estilo para subtitulos
        if 'CustomHeading2' not in [s.name for s in styles]:
            style = styles.add_style('CustomHeading2', WD_STYLE_TYPE.PARAGRAPH)
            style.font.name = 'Arial'
            style.font.size = Pt(14)
            style.font.bold = True
            style.font.color.rgb = self.colors['secondary']
            style.paragraph_format.space_after = Pt(8)
            style.paragraph_format.space_before = Pt(14)

        # Estilo para texto normal
        if 'CustomNormal' not in [s.name for s in styles]:
            style = styles.add_style('CustomNormal', WD_STYLE_TYPE.PARAGRAPH)
            style.font.name = 'Arial'
            style.font.size = Pt(11)
            style.font.color.rgb = self.colors['text']
            style.paragraph_format.space_after = Pt(6)
            style.paragraph_format.line_spacing = 1.15

        # Estilo para metadatos
        if 'CustomMeta' not in [s.name for s in styles]:
            style = styles.add_style('CustomMeta', WD_STYLE_TYPE.PARAGRAPH)
            style.font.name = 'Arial'
            style.font.size = Pt(9)
            style.font.color.rgb = self.colors['muted']

    def _add_header(self, document):
        """Agrega header al documento"""
        section = document.sections[0]
        header = section.header

        header_para = header.paragraphs[0]
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        empresa_info = self._get_empresa_info()

        # Agregar logo si existe
        if self.logo_path:
            run = header_para.add_run()
            run.add_picture(self.logo_path, width=Inches(1.5))
            header_para.add_run('\t\t')

        # Agregar informacion de empresa
        run = header_para.add_run(f"{empresa_info['razon_social']}\n")
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = self.colors['primary']

        if empresa_info['nit']:
            run = header_para.add_run(f"NIT: {empresa_info['nit']}")
            run.font.size = Pt(9)
            run.font.color.rgb = self.colors['muted']

    def _add_footer(self, document, version='1.0'):
        """Agrega footer al documento"""
        section = document.sections[0]
        footer = section.footer

        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = footer_para.add_run(f"Version {version} | Generado: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        run.font.size = Pt(8)
        run.font.color.rgb = self.colors['muted']

        # Agregar numero de pagina
        self._add_page_number(footer_para)

    def _add_page_number(self, paragraph):
        """Agrega numero de pagina al parrafo"""
        paragraph.add_run(' | Pagina ')

        fldSimple = OxmlElement('w:fldSimple')
        fldSimple.set(qn('w:instr'), 'PAGE')

        run = paragraph.add_run()
        run._r.append(fldSimple)

        paragraph.add_run(' de ')

        fldSimple2 = OxmlElement('w:fldSimple')
        fldSimple2.set(qn('w:instr'), 'NUMPAGES')

        run2 = paragraph.add_run()
        run2._r.append(fldSimple2)

    def _add_metadata_table(self, document, metadata):
        """
        Agrega tabla de metadatos del documento.

        Args:
            document: Documento Word
            metadata: Dict con metadatos {label: value}
        """
        # Crear tabla con bordes sutiles
        table = document.add_table(rows=len(metadata), cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Configurar ancho de columnas
        for row_idx, (label, value) in enumerate(metadata.items()):
            row = table.rows[row_idx]

            # Celda de label
            cell_label = row.cells[0]
            cell_label.width = Inches(2)
            para = cell_label.paragraphs[0]
            run = para.add_run(label)
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = self.colors['muted']

            # Celda de valor
            cell_value = row.cells[1]
            cell_value.width = Inches(4)
            para = cell_value.paragraphs[0]
            run = para.add_run(str(value))
            run.font.size = Pt(10)
            run.font.color.rgb = self.colors['text']

        # Agregar shading a las celdas de label
        for row in table.rows:
            cell = row.cells[0]
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'F7FAFC')
            cell._tc.get_or_add_tcPr().append(shading)

        document.add_paragraph()

    def _get_header_color_hex(self):
        """Obtiene el color primario en formato hex para headers de tabla"""
        if self.empresa:
            try:
                from apps.core.models import TenantBranding
                branding = TenantBranding.objects.filter(
                    empresa=self.empresa,
                    is_active=True
                ).first()
                if branding and branding.primary_color:
                    return branding.primary_color.lstrip('#')
            except Exception:
                pass
        return '2C5282'  # Default azul oscuro

    def generate_identidad_completa_docx(self, identity, include_valores=True, include_alcances=True):
        """
        Genera DOCX completo de Identidad Corporativa.

        Args:
            identity: Instancia de CorporateIdentity
            include_valores: Incluir valores
            include_alcances: Incluir alcances

        Returns:
            BytesIO: Buffer con el DOCX generado
        """
        document = Document()
        self._setup_styles(document)
        self._add_header(document)
        self._add_footer(document, version=identity.version)

        empresa_info = self._get_empresa_info()

        # Portada
        para = document.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for _ in range(3):
            para.add_run('\n')

        para = document.add_paragraph('IDENTIDAD CORPORATIVA', style='CustomHeading1')
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        para = document.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(empresa_info['razon_social'])
        run.font.size = Pt(16)
        run.font.color.rgb = self.colors['secondary']

        para = document.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(f"Version {identity.version}\n")
        run.font.size = Pt(12)
        run = para.add_run(f"Vigente desde: {identity.effective_date.strftime('%d/%m/%Y')}")
        run.font.size = Pt(10)
        run.font.color.rgb = self.colors['muted']

        # Nueva pagina - Mision
        document.add_page_break()
        document.add_paragraph('Mision', style='CustomHeading1')

        for paragraph_text in identity.mission.split('\n'):
            if paragraph_text.strip():
                para = document.add_paragraph(paragraph_text, style='CustomNormal')
                para.paragraph_format.first_line_indent = Inches(0.5)

        # Vision
        document.add_paragraph('Vision', style='CustomHeading1')

        for paragraph_text in identity.vision.split('\n'):
            if paragraph_text.strip():
                para = document.add_paragraph(paragraph_text, style='CustomNormal')
                para.paragraph_format.first_line_indent = Inches(0.5)

        # Valores Corporativos
        if include_valores:
            valores = identity.values.filter(is_active=True).order_by('orden')
            if valores.exists():
                document.add_page_break()
                document.add_paragraph('Valores Corporativos', style='CustomHeading1')

                for valor in valores:
                    # Nombre del valor
                    para = document.add_paragraph()
                    run = para.add_run(f'\u2022 {valor.name}')
                    run.bold = True
                    run.font.size = Pt(12)
                    run.font.color.rgb = self.colors['secondary']

                    # Descripcion
                    para = document.add_paragraph(valor.description, style='CustomNormal')
                    para.paragraph_format.left_indent = Inches(0.3)

        # Alcances del Sistema
        if include_alcances:
            alcances = identity.alcances.filter(is_active=True)
            if alcances.exists():
                document.add_page_break()
                document.add_paragraph('Alcance del Sistema de Gestion', style='CustomHeading1')

                # Tabla de alcances
                table = document.add_table(rows=1, cols=4)
                table.style = 'Table Grid'

                # Headers
                header_color = self._get_header_color_hex()
                headers = ['Norma ISO', 'Alcance', 'Estado', 'Organismo']
                header_row = table.rows[0]
                for idx, header in enumerate(headers):
                    cell = header_row.cells[idx]
                    para = cell.paragraphs[0]
                    run = para.add_run(header)
                    run.bold = True
                    run.font.size = Pt(9)

                    shading = OxmlElement('w:shd')
                    shading.set(qn('w:fill'), header_color)
                    cell._tc.get_or_add_tcPr().append(shading)
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

                # Datos
                for alcance in alcances:
                    row = table.add_row()
                    row.cells[0].text = alcance.norma_iso.short_name if alcance.norma_iso else 'N/A'
                    row.cells[1].text = alcance.scope[:80] + '...' if len(alcance.scope) > 80 else alcance.scope
                    row.cells[2].text = 'Certificado' if alcance.is_certified else 'No Certificado'
                    row.cells[3].text = alcance.certification_body or 'N/A'

        # Control de documento
        document.add_page_break()
        document.add_paragraph('Control de Documento', style='CustomHeading1')

        control_metadata = {
            'Fecha de generacion': datetime.now().strftime('%d/%m/%Y %H:%M'),
            'Creado por': identity.created_by.get_full_name() if identity.created_by else 'Sistema',
            'Ultima actualizacion': identity.updated_at.strftime('%d/%m/%Y %H:%M') if identity.updated_at else 'N/A',
            'Version del documento': identity.version,
        }
        self._add_metadata_table(document, control_metadata)

        # Guardar a buffer
        docx_buffer = BytesIO()
        document.save(docx_buffer)
        docx_buffer.seek(0)

        return docx_buffer


# Funcion de conveniencia
def generar_docx_identidad_completa(identity, empresa=None, **options):
    """
    Genera DOCX de identidad corporativa completa.

    Args:
        identity: Instancia de CorporateIdentity
        empresa: Instancia de EmpresaConfig (opcional)
        **options: include_valores, include_alcances

    Returns:
        BytesIO: Buffer con DOCX
    """
    generator = IdentidadDOCXGenerator(empresa=empresa)
    return generator.generate_identidad_completa_docx(identity, **options)
