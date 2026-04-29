"""
Generador de DOCX para Gestion Documental
==========================================
Genera documentos Word editables de documentos del sistema de gestion
usando python-docx.

Patron: Identico a identidad/exporters/docx_generator.py
"""
import os
from io import BytesIO
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def hex_to_rgb_color(hex_color: str):
    hex_color = hex_color.lstrip('#')
    try:
        r = int(hex_color[0:2], 16)
        g = int(hex_color[2:4], 16)
        b = int(hex_color[4:6], 16)
        return RGBColor(r, g, b)
    except (ValueError, IndexError):
        return RGBColor(0x2d, 0x37, 0x48)


class DocumentoDOCXGenerator:
    """Genera documentos DOCX de documentos del sistema de gestion documental."""

    def __init__(self, empresa=None):
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx no esta instalado. Ejecute: pip install python-docx")

        self.empresa = empresa
        self.logo_path = None
        self.colors = self._load_colors()
        if empresa and hasattr(empresa, 'logo') and empresa.logo:
            self._load_logo()

    def _load_colors(self):
        defaults = {
            'primary': '#1a365d',
            'secondary': '#3182ce',
            'text': '#2d3748',
            'muted': '#718096',
            'success': '#48bb78',
        }
        if self.empresa:
            try:
                from apps.core.models import TenantBranding
                branding = TenantBranding.objects.filter(
                    empresa=self.empresa, is_active=True
                ).first()
                if branding:
                    defaults['primary'] = branding.primary_color or defaults['primary']
                    defaults['secondary'] = branding.secondary_color or defaults['secondary']
            except Exception:
                pass

        return {key: hex_to_rgb_color(value) for key, value in defaults.items()}

    def _load_logo(self):
        try:
            if self.empresa.logo and hasattr(self.empresa.logo, 'path'):
                if os.path.exists(self.empresa.logo.path):
                    self.logo_path = self.empresa.logo.path
        except Exception:
            pass

    def _get_empresa_info(self):
        if self.empresa:
            return {
                'razon_social': self.empresa.razon_social or 'Empresa',
                'nit': self.empresa.nit or '',
            }
        return {'razon_social': 'Empresa', 'nit': ''}

    def _setup_styles(self, document):
        styles = document.styles

        if 'CustomHeading1' not in [s.name for s in styles]:
            style = styles.add_style('CustomHeading1', WD_STYLE_TYPE.PARAGRAPH)
            style.font.name = 'Arial'
            style.font.size = Pt(18)
            style.font.bold = True
            style.font.color.rgb = self.colors['primary']
            style.paragraph_format.space_after = Pt(12)
            style.paragraph_format.space_before = Pt(18)

        if 'CustomHeading2' not in [s.name for s in styles]:
            style = styles.add_style('CustomHeading2', WD_STYLE_TYPE.PARAGRAPH)
            style.font.name = 'Arial'
            style.font.size = Pt(14)
            style.font.bold = True
            style.font.color.rgb = self.colors['secondary']
            style.paragraph_format.space_after = Pt(8)
            style.paragraph_format.space_before = Pt(14)

        if 'CustomNormal' not in [s.name for s in styles]:
            style = styles.add_style('CustomNormal', WD_STYLE_TYPE.PARAGRAPH)
            style.font.name = 'Arial'
            style.font.size = Pt(11)
            style.font.color.rgb = self.colors['text']
            style.paragraph_format.space_after = Pt(6)
            style.paragraph_format.line_spacing = 1.15

    def _add_header(self, document, codigo=''):
        section = document.sections[0]
        header = section.header
        header_para = header.paragraphs[0]
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        empresa_info = self._get_empresa_info()

        if self.logo_path:
            run = header_para.add_run()
            run.add_picture(self.logo_path, width=Inches(1.5))
            header_para.add_run('\t\t')

        run = header_para.add_run(f"{empresa_info['razon_social']}\n")
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = self.colors['primary']

        if empresa_info['nit']:
            run = header_para.add_run(f"NIT: {empresa_info['nit']}")
            run.font.size = Pt(9)
            run.font.color.rgb = self.colors['muted']

        if codigo:
            header_para.add_run('\n')
            run = header_para.add_run(codigo)
            run.font.size = Pt(9)
            run.font.color.rgb = self.colors['muted']

    def _add_footer(self, document, codigo='', version='1.0'):
        section = document.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = footer_para.add_run(
            f"{codigo} | v{version} | Generado: {datetime.now().strftime('%d/%m/%Y %H:%M')}"
        )
        run.font.size = Pt(8)
        run.font.color.rgb = self.colors['muted']

        footer_para.add_run(' | Pagina ')
        fld = OxmlElement('w:fldSimple')
        fld.set(qn('w:instr'), 'PAGE')
        run = footer_para.add_run()
        run._r.append(fld)
        footer_para.add_run(' de ')
        fld2 = OxmlElement('w:fldSimple')
        fld2.set(qn('w:instr'), 'NUMPAGES')
        run2 = footer_para.add_run()
        run2._r.append(fld2)

    def _add_metadata_table(self, document, metadata):
        table = document.add_table(rows=len(metadata), cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for row_idx, (label, value) in enumerate(metadata.items()):
            row = table.rows[row_idx]

            cell_label = row.cells[0]
            cell_label.width = Inches(2)
            para = cell_label.paragraphs[0]
            run = para.add_run(label)
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = self.colors['muted']

            cell_value = row.cells[1]
            cell_value.width = Inches(4)
            para = cell_value.paragraphs[0]
            run = para.add_run(str(value or 'N/A'))
            run.font.size = Pt(10)
            run.font.color.rgb = self.colors['text']

        for row in table.rows:
            cell = row.cells[0]
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'F7FAFC')
            cell._tc.get_or_add_tcPr().append(shading)

        document.add_paragraph()

    def generate_documento_docx(self, documento):
        """
        Genera DOCX de un documento del sistema de gestion.

        Args:
            documento: Instancia de Documento

        Returns:
            BytesIO: Buffer con el DOCX generado
        """
        doc = Document()
        self._setup_styles(doc)
        self._add_header(doc, codigo=documento.codigo)
        self._add_footer(doc, codigo=documento.codigo, version=documento.version_actual)

        # Titulo
        para = doc.add_paragraph(documento.titulo, style='CustomHeading1')
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata
        metadata = {
            'Codigo': documento.codigo,
            'Version': documento.version_actual,
            'Estado': documento.get_estado_display(),
            'Clasificacion': documento.get_clasificacion_display(),
            'Tipo de Documento': documento.tipo_documento.nombre,
            'Elaborado por': documento.elaborado_por.get_full_name(),
            'Revisado por': (
                documento.revisado_por.get_full_name()
                if documento.revisado_por else 'N/A'
            ),
            'Aprobado por': (
                documento.aprobado_por.get_full_name()
                if documento.aprobado_por else 'N/A'
            ),
            'Fecha Publicacion': (
                documento.fecha_publicacion.strftime('%d/%m/%Y')
                if documento.fecha_publicacion else 'N/A'
            ),
            'Fecha Vigencia': (
                documento.fecha_vigencia.strftime('%d/%m/%Y')
                if documento.fecha_vigencia else 'N/A'
            ),
        }
        self._add_metadata_table(doc, metadata)

        # Contenido
        doc.add_paragraph('Contenido', style='CustomHeading2')
        contenido = self._sustituir_variables(documento)

        # Convertir HTML basico a texto plano para DOCX
        contenido_texto = self._html_to_text(contenido)
        for paragraph_text in contenido_texto.split('\n'):
            if paragraph_text.strip():
                doc.add_paragraph(paragraph_text.strip(), style='CustomNormal')

        # Control de documento
        doc.add_page_break()
        doc.add_paragraph('Control de Documento', style='CustomHeading1')
        control = {
            'Fecha de generacion': datetime.now().strftime('%d/%m/%Y %H:%M'),
            'Numero de revision': str(documento.numero_revision),
            'Revision programada': (
                documento.fecha_revision_programada.strftime('%d/%m/%Y')
                if documento.fecha_revision_programada else 'N/A'
            ),
        }
        self._add_metadata_table(doc, control)

        # Guardar
        docx_buffer = BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)
        return docx_buffer

    def _sustituir_variables(self, documento):
        """Sustituye variables {{var}} en el contenido del documento."""
        contenido = documento.contenido or ''
        empresa_info = self._get_empresa_info()

        variables = {
            '{{codigo}}': documento.codigo or '',
            '{{titulo}}': documento.titulo or '',
            '{{version}}': documento.version_actual or '',
            '{{estado}}': documento.get_estado_display(),
            '{{empresa}}': empresa_info['razon_social'],
            '{{nit}}': empresa_info['nit'],
            '{{fecha_publicacion}}': (
                documento.fecha_publicacion.strftime('%d/%m/%Y')
                if documento.fecha_publicacion else ''
            ),
            '{{fecha_vigencia}}': (
                documento.fecha_vigencia.strftime('%d/%m/%Y')
                if documento.fecha_vigencia else ''
            ),
            '{{elaborado_por}}': (
                documento.elaborado_por.get_full_name()
                if documento.elaborado_por else ''
            ),
            '{{revisado_por}}': (
                documento.revisado_por.get_full_name()
                if documento.revisado_por else ''
            ),
            '{{aprobado_por}}': (
                documento.aprobado_por.get_full_name()
                if documento.aprobado_por else ''
            ),
        }

        for var, value in variables.items():
            contenido = contenido.replace(var, value)

        return contenido

    @staticmethod
    def _html_to_text(html_content):
        """Convierte HTML basico a texto plano para DOCX."""
        import re
        text = html_content
        # Reemplazar <br>, <p>, <div> con saltos de linea
        text = re.sub(r'<br\s*/?>',  '\n', text)
        text = re.sub(r'</p>',  '\n', text)
        text = re.sub(r'</div>',  '\n', text)
        text = re.sub(r'</li>',  '\n', text)
        text = re.sub(r'<li[^>]*>',  '  - ', text)
        # Remover tags HTML restantes
        text = re.sub(r'<[^>]+>', '', text)
        # Decodificar entidades HTML basicas
        text = text.replace('&amp;', '&')
        text = text.replace('&lt;', '<')
        text = text.replace('&gt;', '>')
        text = text.replace('&nbsp;', ' ')
        text = text.replace('&quot;', '"')
        # Limpiar lineas vacias multiples
        text = re.sub(r'\n{3,}', '\n\n', text)
        return text.strip()
